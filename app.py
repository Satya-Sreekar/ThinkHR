import os
import locale
import pandas as pd
from datetime import datetime, timedelta
from flask import Flask, render_template, request, redirect, url_for, flash, send_file
from flask_mysqldb import MySQL
from flask_login import LoginManager, login_user, logout_user, login_required, current_user, UserMixin
from docx import Document
from docx.shared import Pt
import docx2pdf
from io import BytesIO
import yaml
from werkzeug.utils import secure_filename
import sys
import subprocess
if sys.platform.startswith('win'):
    import pythoncom

# --------------------------------------------------------------------------
# App Initialization and Configuration
# --------------------------------------------------------------------------
app = Flask(__name__)
app.secret_key = 'your_secret_key'
app.config['UPLOAD_FOLDER'] = os.path.join('static', 'uploads')

# Load database configuration from YAML file
with open('db.yaml', 'r') as db_file:
    db_config = yaml.safe_load(db_file)
app.config['MYSQL_HOST'] = db_config['mysql_host']
app.config['MYSQL_USER'] = db_config['mysql_user']
app.config['MYSQL_PASSWORD'] = db_config['mysql_password']
app.config['MYSQL_DB'] = db_config['mysql_db']

mysql = MySQL(app)

# Initialize Flask-Login
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'

# --------------------------------------------------------------------------
# Locale and Formatting Setup
# --------------------------------------------------------------------------
locale.setlocale(locale.LC_ALL, 'en_IN.UTF-8')

def format_number_indian(number):
    """Format a number in the Indian numbering system with commas and conditional paise."""
    integer_part, _, decimal_part = f"{number:.2f}".partition(".")
    integer_formatted = str(integer_part)[::-1]
    formatted = []
    for i, digit in enumerate(integer_formatted):
        if i > 2 and (i - 3) % 2 == 0:
            formatted.append(',')
        formatted.append(digit)
    integer_with_commas = ''.join(formatted[::-1])
    if decimal_part != "00":
        return f"{integer_with_commas}.{decimal_part}"
    return integer_with_commas

def format_category(category):
    """Capitalize the first letter of each word in the category."""
    return ' '.join(word.capitalize() for word in category.split())

# --------------------------------------------------------------------------
# User Model and Loader
# --------------------------------------------------------------------------
class User(UserMixin):
    def __init__(self, id, username, role):
        self.id = id
        self.username = username
        self.role = role

@login_manager.user_loader
def load_user(user_id):
    cursor = mysql.connection.cursor()
    cursor.execute("SELECT id, username, role FROM staff WHERE id = %s", [user_id])
    user = cursor.fetchone()
    if user:
        return User(id=user[0], username=user[1], role=user[2])
    return None

# --------------------------------------------------------------------------
# Dashboard Route (Merged with Financial Overview)
# --------------------------------------------------------------------------
@app.route('/')
@login_required
def index():
    cur = mysql.connection.cursor()
    # Staff/attendance data
    cur.execute("SELECT COUNT(*) FROM staff")
    total_staff = cur.fetchone()[0]
    cur.execute("SELECT COUNT(DISTINCT staff_id) FROM attendance WHERE DATE(checkin_time) = CURDATE()")
    checked_in_today = cur.fetchone()[0]
    cur.execute("SELECT COUNT(id) FROM department")
    department_count = cur.fetchone()[0]
    return render_template(
        'index.html', 
        total_staff=total_staff, 
        present_today=checked_in_today, 
        department_count=department_count,
    )

# --------------------------------------------------------------------------
# Staff Management Routes
# --------------------------------------------------------------------------
@app.route('/ThinkStaff')
@login_required
def ThinkStaff():
    cur = mysql.connection.cursor()
    cur.execute("SELECT COUNT(*) FROM staff")
    total_staff = cur.fetchone()[0]
    cur.execute("SELECT COUNT(DISTINCT staff_id) FROM attendance WHERE DATE(checkin_time) = CURDATE()")
    checked_in_today = cur.fetchone()[0]
    cur.execute("SELECT COUNT(id) FROM department")
    department_count = cur.fetchone()[0]
    return render_template('ThinkStaff.html',
                           total_staff=total_staff, 
                           present_today=checked_in_today,
                           department_count=department_count)

@app.route('/mannage_staff')
@login_required
def mannage_staff():
    cursor = mysql.connection.cursor()
    cursor.execute("""
        SELECT s.id, s.name, s.position, d.name AS department
        FROM staff s LEFT JOIN department d ON s.department = d.id
    """)
    staff_list = cursor.fetchall()
    return render_template('mannage_staff.html', staff_list=staff_list)

@app.route('/onboard', methods=['GET'])
@login_required
def onboard():
    cursor = mysql.connection.cursor()
    cursor.execute("SELECT id, name FROM staff")
    staff_list = cursor.fetchall()
    cursor.execute("SELECT id, name FROM department")
    department_list = cursor.fetchall()
    return render_template('onboard.html', staff_list=staff_list, department_list=department_list)

@app.route('/add_staff', methods=['POST'])
@login_required
def add_staff():
    name = request.form['name']
    department = request.form['department']
    reportee = request.form['reportee']
    email = request.form['email']
    positon = request.form['position']
    if name and department and reportee and email:
        cursor = mysql.connection.cursor()
        cursor.execute(
            "INSERT INTO staff (name, department, reportee, email, position) VALUES (%s, %s, %s, %s, %s)",
            (name, department, reportee, email, positon)
        )
        mysql.connection.commit()
        flash('Staff added successfully!', 'success')
    else:
        flash('Please fill in all fields.', 'danger')
    return redirect(url_for('mannage_staff'))

@app.route('/edit_staff/<int:staff_id>', methods=['GET', 'POST'])
@login_required
def edit_staff(staff_id):
    cursor = mysql.connection.cursor()
    cursor.execute("SELECT id, name FROM staff")
    staff_list = cursor.fetchall()
    cursor.execute("SELECT id, name FROM department")
    department_list = cursor.fetchall()
    if request.method == 'POST':
        name = request.form['name']
        department = request.form['department']
        email = request.form['email']
        reportee = request.form['reportee']
        position = request.form['position']
        cursor.execute("""
            UPDATE staff SET name = %s, department = %s, email = %s, reportee = %s, position = %s 
            WHERE id = %s
        """, (name, department, email, reportee, position, staff_id))
        mysql.connection.commit()
        flash('Staff updated successfully!', 'success')
        return redirect(url_for('mannage_staff'))
    cursor.execute("SELECT * FROM staff WHERE id = %s", [staff_id])
    staff = cursor.fetchone()
    return render_template('edit_staff.html', staff=staff, staff_list=staff_list, department_list=department_list)

@app.route('/delete_staff/<int:staff_id>')
@login_required
def delete_staff(staff_id):
    cursor = mysql.connection.cursor()
    cursor.execute("DELETE FROM staff WHERE id = %s", [staff_id])
    mysql.connection.commit()
    flash('Staff deleted successfully!', 'success')
    return redirect(url_for('mannage_staff'))

# --------------------------------------------------------------------------
# Department Management Routes
# --------------------------------------------------------------------------
@app.route('/departments')
@login_required
def department():
    cursor = mysql.connection.cursor()
    cursor.execute("SELECT * FROM department")
    department_list = cursor.fetchall()
    return render_template('departments.html', department_list=department_list)

@app.route('/add_department', methods=['POST'])
@login_required
def add_department():
    name = request.form['department_name']
    if name:
        cursor = mysql.connection.cursor()
        cursor.execute("INSERT INTO department (name) VALUES (%s)", [name])
        mysql.connection.commit()
        flash('Department added successfully!', 'success')
    else:
        flash('Please fill in all fields.', 'danger')
    return redirect(url_for('department'))

# --------------------------------------------------------------------------
# Attendance Routes
# --------------------------------------------------------------------------
@app.route('/attendence')
@login_required
def attendence():
    cursor = mysql.connection.cursor()
    cursor.execute("""
        SELECT s.id, s.name, d.name AS department,
          CASE 
            WHEN (SELECT COUNT(*) FROM attendance a WHERE a.staff_id = s.id AND a.checkout_time IS NULL) > 0 
            THEN 'checked_in'
            ELSE 'checked_out'
          END AS attendance_status
        FROM staff s
        LEFT JOIN department d ON s.department = d.id
    """)
    staff_list = cursor.fetchall()
    return render_template('attendence.html', staff_list=staff_list)

@app.route('/checkin/<int:staff_id>', methods=['POST'])
@login_required
def checkin(staff_id):
    cursor = mysql.connection.cursor()
    cursor.execute("INSERT INTO attendance (staff_id, checkin_time) VALUES (%s, NOW())", [staff_id])
    mysql.connection.commit()
    flash('Staff checked in successfully!', 'success')
    return redirect(url_for('attendence'))

@app.route('/checkout/<int:staff_id>', methods=['POST'])
@login_required
def checkout(staff_id):
    cursor = mysql.connection.cursor()
    cursor.execute("""
        UPDATE attendance 
        SET checkout_time = NOW() 
        WHERE staff_id = %s AND checkout_time IS NULL
        ORDER BY checkin_time DESC LIMIT 1
    """, [staff_id])
    mysql.connection.commit()
    flash('Staff checked out successfully!', 'success')
    return redirect(url_for('attendence'))

# --------------------------------------------------------------------------
# Attendance Summary Routes (Daily/Weekly)
# --------------------------------------------------------------------------
@app.route('/summary')
@login_required
def summary():
    period = request.args.get('period', 'daily')
    selected_date_str = request.args.get('date')
    try:
        selected_date = datetime.strptime(selected_date_str, '%Y-%m-%d').date() if selected_date_str else datetime.now().date()
    except (ValueError, TypeError):
        selected_date = datetime.now().date()
    cursor = mysql.connection.cursor()
    cursor.execute("SELECT MIN(DATE(checkin_time)), MAX(DATE(checkin_time)) FROM attendance")
    min_date, max_date = cursor.fetchone()
    if not min_date or not max_date:
        min_date = max_date = None
    if period == 'daily':
        cursor.execute("""
            SELECT staff.name, 
                   MIN(attendance.checkin_time) AS first_checkin, 
                   MAX(attendance.checkout_time) AS last_checkout,
                   TIMEDIFF(MAX(attendance.checkout_time), MIN(attendance.checkin_time)) AS effective_hours
            FROM attendance 
            JOIN staff ON attendance.staff_id = staff.id
            WHERE DATE(attendance.checkin_time) = %s
            GROUP BY staff.id
        """, [selected_date])
        summary_data = cursor.fetchall()
        previous_date = (selected_date - timedelta(days=1)) if min_date and selected_date > min_date else None
        next_date = (selected_date + timedelta(days=1)) if max_date and selected_date < max_date and selected_date < datetime.now().date() else None
        return render_template(
            'summary.html',
            summary_data=summary_data,
            period='daily',
            selected_date=selected_date,
            previous_date=previous_date,
            next_date=next_date,
            min_date=min_date,
            max_date=max_date,
            timedelta=timedelta
        )
    elif period == 'weekly':
        week_start = selected_date - timedelta(days=selected_date.weekday())
        week_end = week_start + timedelta(days=6)
        cursor.execute("""
            SELECT staff.name, DATE(attendance.checkin_time) AS day, 
                   MIN(attendance.checkin_time) AS first_checkin, 
                   MAX(attendance.checkout_time) AS last_checkout,
                   TIMEDIFF(MAX(attendance.checkout_time), MIN(attendance.checkin_time)) AS effective_hours
            FROM attendance 
            JOIN staff ON attendance.staff_id = staff.id
            WHERE DATE(attendance.checkin_time) BETWEEN %s AND %s
            GROUP BY staff.id, day
        """, (week_start, week_end))
        weekly_data = cursor.fetchall()
        summary_data = {}
        for staff_name, day, first_checkin, last_checkout, effective_hours in weekly_data:
            if staff_name not in summary_data:
                summary_data[staff_name] = {
                    week_start + timedelta(days=i): {
                        'status': 'absent', 
                        'first_checkin': None, 
                        'last_checkout': None, 
                        'effective_hours': None
                    } for i in range(7)
                }
            summary_data[staff_name][day] = {
                'status': 'present', 
                'first_checkin': first_checkin, 
                'last_checkout': last_checkout, 
                'effective_hours': effective_hours
            }
        previous_week = (week_start - timedelta(weeks=1)) if min_date and week_start > min_date else None
        next_week = (week_start + timedelta(weeks=1)) if max_date and week_start + timedelta(weeks=1) <= datetime.now().date() else None
        return render_template(
            'summary.html',
            summary_data=summary_data,
            period='weekly',
            week_start=week_start,
            week_end=week_end,
            previous_week=previous_week,
            next_week=next_week,
            min_date=min_date,
            max_date=max_date,
            timedelta=timedelta
        )

# --------------------------------------------------------------------------
# Company Hierarchy Route
# --------------------------------------------------------------------------
@app.route('/company_hierarchy')
@login_required
def company_hierarchy():
    cursor = mysql.connection.cursor()
    cursor.execute("""
        SELECT s.id, s.name, s.reportee, d.name AS department_name, s.position
        FROM staff s
        JOIN department d ON s.department = d.id
        ORDER BY d.name ASC, s.reportee ASC
    """)
    staff_data = cursor.fetchall()
    staff_dict = {
        staff[0]: {
            'id': staff[0],
            'name': staff[1],
            'reportee': staff[2],
            'department': staff[3],
            'position': staff[4],
            'subordinates': []
        }
        for staff in staff_data
    }
    for staff in staff_data:
        staff_id, name, manager_id, department_name, position = staff
        if manager_id != 0 and manager_id in staff_dict:
            staff_dict[manager_id]['subordinates'].append(staff_dict[staff_id])
    top_managers = [staff_dict[staff[0]] for staff in staff_data if staff[2] == 0]
    def bfs_tree(manager):
        from collections import deque
        queue = deque([(manager, 0)])
        tree = []
        while queue:
            node, level = queue.popleft()
            if len(tree) <= level:
                tree.append([])
            tree[level].append({
                'name': node['name'],
                'department': node['department'],
                'position': node['position']
            })
            for sub in sorted(node['subordinates'], key=lambda x: x['id']):
                queue.append((sub, level + 1))
        return tree
    trees = [bfs_tree(manager) for manager in top_managers]
    return render_template('hierarchy.html', trees=trees)

# --------------------------------------------------------------------------
# Invoice Generation Routes and Functions
# --------------------------------------------------------------------------
def generate_invoice(data):
    doc = Document("invoice.docx")
    for para in doc.paragraphs:
        if "Grand Total:" in para.text:
            para.text = f"Grand Total: Rupees {data['grand_total_text']} Rs.{data['grand_total']}/-"
    table1 = doc.tables[0]
    invoice_date = datetime.strptime(data["invoice_date"], "%Y-%m-%d").strftime("%B %d, %Y")
    table1.rows[1].cells[0].text = invoice_date
    run = table1.rows[1].cells[0].paragraphs[0].runs[0]
    run.font.size = Pt(16)
    table1.rows[1].cells[1].text = f"INVOICE#{data['invoice_number']}"
    run = table1.rows[1].cells[1].paragraphs[0].runs[0]
    run.font.size = Pt(16)
    table2 = doc.tables[1]
    invoice_date_obj = datetime.strptime(data["invoice_date"], "%Y-%m-%d")
    due_date = (invoice_date_obj + timedelta(days=14)).strftime("%B %d, %Y")
    table2.rows[1].cells[3].text = due_date
    table3 = doc.tables[2]
    new_row = table3.add_row()
    new_row.cells[0].text = "1"
    new_row.cells[1].text = data["billable_hours"]
    new_row.cells[2].text = f"Rs. {data['amount_per_hour']}"
    new_row.cells[3].text = data["start_date"]
    new_row.cells[4].text = data["end_date"]
    new_row.cells[5].text = f"Rs. {data['subtotal']}"
    table5 = doc.tables[4]
    table5.rows[0].cells[1].text = f"Rs. {data['subtotal']}"
    table5.rows[1].cells[1].text = f"Rs. {data['gst']}"
    table5.rows[2].cells[1].text = f"Rs. {data['grand_total']}"
    output_stream = BytesIO()
    doc.save(output_stream)
    output_stream.seek(0)
    return output_stream

@app.route('/ThinkInvoice')
@login_required
def ThinkInvoice():
    return render_template('ThinkInvoice.html')

@app.route('/reciept')
@login_required
def reciept():
    return render_template('invoice_form.html')

@app.route('/generate', methods=['POST'])
@login_required
def pdf_download():
    file_name = generate()
    try:
        if sys.platform.startswith('win'):
            # Windows-specific conversion using docx2pdf
            pythoncom.CoInitialize()
            docx2pdf.convert(file_name)
            pythoncom.CoUninitialize()
        else:
            # Non-Windows: use LibreOffice headless conversion
            subprocess.run([
                'libreoffice', '--headless', '--convert-to', 'pdf', file_name,
                '--outdir', os.path.dirname(file_name)
            ], check=True)
    except Exception as e:
        flash(f"Error converting file: {e}", 'danger')
        return redirect(url_for('reciept'))

    pdf_file_name = file_name.replace("docx", "pdf")
    return send_file(pdf_file_name, as_attachment=True)

def generate():
    form_data = request.form
    invoice_data = {
        "invoice_date": form_data.get("invoice_date"),
        "invoice_number": form_data.get("invoice_number"),
        "billable_hours": form_data.get("billable_hours"),
        "amount_per_hour": form_data.get("amount_per_hour"),
        "start_date": form_data.get("start_date"),
        "end_date": form_data.get("end_date"),
        "subtotal": form_data.get("subtotal"),
        "gst": form_data.get("gst"),
        "grand_total": form_data.get("grand_total"),
        "grand_total_text": form_data.get("grand_total_text"),
        "total_amount": form_data.get("total_amount"),
    }
    invoice_file = generate_invoice(invoice_data)
    file_name = f"invoices/invoice_{invoice_data['invoice_number']}.docx"
    with open(file_name, "wb") as f:
        f.write(invoice_file.getvalue())
    return file_name


# --------------------------------------------------------------------------
# Financial Transactions and Related Routes
# --------------------------------------------------------------------------
@app.route('/ThinkTransactions')
@login_required
def transactions():
    cur = mysql.connection.cursor()

    # Total financial data
    cur.execute("SELECT IFNULL(SUM(amount), 0) FROM transactions WHERE type = 'Income'")
    total_income_val = cur.fetchone()[0]
    cur.execute("SELECT IFNULL(SUM(amount), 0) FROM transactions WHERE type = 'Expense'")
    total_expense_val = cur.fetchone()[0]
    balance_val = total_income_val - total_expense_val

    total_income = format_number_indian(total_income_val)
    total_expense = format_number_indian(total_expense_val)
    balance = format_number_indian(balance_val)

    # Current Year
    current_year = datetime.now().year
    cur.execute(f"SELECT IFNULL(SUM(amount), 0) FROM transactions WHERE type = 'Income' AND YEAR(date) = {current_year}")
    current_year_income_val = cur.fetchone()[0]
    cur.execute(f"SELECT IFNULL(SUM(amount), 0) FROM transactions WHERE type = 'Expense' AND YEAR(date) = {current_year}")
    current_year_expense_val = cur.fetchone()[0]
    current_year_balance_val = current_year_income_val - current_year_expense_val

    current_year_income = format_number_indian(current_year_income_val)
    current_year_expense = format_number_indian(current_year_expense_val)
    current_year_balance = format_number_indian(current_year_balance_val)

    # Previous Year
    previous_year = current_year - 1
    cur.execute(f"SELECT IFNULL(SUM(amount), 0) FROM transactions WHERE type = 'Income' AND YEAR(date) = {previous_year}")
    previous_year_income_val = cur.fetchone()[0]
    cur.execute(f"SELECT IFNULL(SUM(amount), 0) FROM transactions WHERE type = 'Expense' AND YEAR(date) = {previous_year}")
    previous_year_expense_val = cur.fetchone()[0]
    previous_year_balance_val = previous_year_income_val - previous_year_expense_val

    previous_year_income = format_number_indian(previous_year_income_val)
    previous_year_expense = format_number_indian(previous_year_expense_val)
    previous_year_balance = format_number_indian(previous_year_balance_val)

    # Current Month
    current_month = datetime.now().month
    cur.execute(f"""
        SELECT IFNULL(SUM(amount), 0) FROM transactions 
        WHERE type = 'Income' AND YEAR(date) = {current_year} AND MONTH(date) = {current_month}
    """)
    current_month_income_val = cur.fetchone()[0]
    cur.execute(f"""
        SELECT IFNULL(SUM(amount), 0) FROM transactions 
        WHERE type = 'Expense' AND YEAR(date) = {current_year} AND MONTH(date) = {current_month}
    """)
    current_month_expense_val = cur.fetchone()[0]
    current_month_balance_val = current_month_income_val - current_month_expense_val

    current_month_income = format_number_indian(current_month_income_val)
    current_month_expense = format_number_indian(current_month_expense_val)
    current_month_balance = format_number_indian(current_month_balance_val)

    # Previous Month
    first_day_current_month = datetime.now().replace(day=1)
    last_day_previous_month = first_day_current_month - timedelta(days=1)
    previous_month_year = last_day_previous_month.year
    previous_month = last_day_previous_month.month

    cur.execute(f"""
        SELECT IFNULL(SUM(amount), 0) FROM transactions 
        WHERE type = 'Income' AND YEAR(date) = {previous_month_year} AND MONTH(date) = {previous_month}
    """)
    previous_month_income_val = cur.fetchone()[0]
    cur.execute(f"""
        SELECT IFNULL(SUM(amount), 0) FROM transactions 
        WHERE type = 'Expense' AND YEAR(date) = {previous_month_year} AND MONTH(date) = {previous_month}
    """)
    previous_month_expense_val = cur.fetchone()[0]
    previous_month_balance_val = previous_month_income_val - previous_month_expense_val

    previous_month_income = format_number_indian(previous_month_income_val)
    previous_month_expense = format_number_indian(previous_month_expense_val)
    previous_month_balance = format_number_indian(previous_month_balance_val)

    cur.close()

    return render_template(
        'ThinkTransactions.html',
        total_income=total_income,
        total_expense=total_expense,
        balance=balance,
        current_year_income=current_year_income,
        current_year_expense=current_year_expense,
        current_year_balance=current_year_balance,
        previous_year_income=previous_year_income,
        previous_year_expense=previous_year_expense,
        previous_year_balance=previous_year_balance,
        current_month_income=current_month_income,
        current_month_expense=current_month_expense,
        current_month_balance=current_month_balance,
        previous_month_income=previous_month_income,
        previous_month_expense=previous_month_expense,
        previous_month_balance=previous_month_balance,
    )


@app.route('/add', methods=["GET", "POST"])
@login_required
def add_entry():
    if request.method == "POST":
        type_ = request.form["type"]
        category = format_category(request.form["category"])
        amount = float(request.form["amount"])
        date = request.form["date"]
        description = request.form.get("description", "")

        cur = mysql.connection.cursor()
        cur.execute("""
            INSERT INTO transactions (type, category, amount, date, description)
            VALUES (%s, %s, %s, %s, %s)
        """, (type_, category, amount, date, description))
        cur.execute("SELECT DISTINCT category FROM transactions WHERE category = %s", (category,))
        if not cur.fetchone():
            cur.execute("INSERT INTO categories (name) VALUES (%s)", (category,))
        mysql.connection.commit()
        cur.close()

        flash('Transaction added successfully!', 'success')
        return redirect(url_for("index"))
    return render_template("add_entry.html")

@app.route('/view')
@login_required
def view_entries():
    page = request.args.get('page', 1, type=int)
    per_page = 10
    offset = (page - 1) * per_page

    sort_by = request.args.get('sort_by', 'date')
    sort_order = request.args.get('sort_order', 'desc')

    valid_columns = ['id', 'type', 'category', 'amount', 'date']
    if sort_by not in valid_columns:
        sort_by = 'date'
    sort_order = 'asc' if sort_order == 'asc' else 'desc'

    cur = mysql.connection.cursor()
    cur.execute("SELECT COUNT(*) FROM transactions")
    total_entries = cur.fetchone()[0]

    total_pages = (total_entries + per_page - 1) // per_page
    if page < 1 or page > total_pages:
        flash('Invalid page number!', 'danger')
        return redirect(url_for('view_entries', page=1))

    cur.execute(f"""
        SELECT * FROM transactions
        ORDER BY {sort_by} {sort_order}
        LIMIT %s OFFSET %s
    """, (per_page, offset))
    transactions = cur.fetchall()
    cur.close()

    formatted_transactions = [
        (
            t[0],
            t[1],
            t[2],
            format_number_indian(t[3]),
            datetime.strptime(str(t[4]), "%Y-%m-%d").strftime("%d-%m-%Y"),
            t[5]
        ) for t in transactions
    ]

    return render_template(
        "view_entries.html",
        transactions=formatted_transactions,
        page=page,
        total_pages=total_pages,
        sort_by=sort_by,
        sort_order=sort_order
    )

@app.route('/get_categories', methods=["GET"])
@login_required
def get_categories():
    cur = mysql.connection.cursor()
    cur.execute("SELECT DISTINCT category FROM transactions")
    categories = [row[0] for row in cur.fetchall()]
    cur.close()
    return {"categories": categories}

@app.route('/edit/<int:transaction_id>', methods=["GET", "POST"])
@login_required
def edit_transaction(transaction_id):
    cur = mysql.connection.cursor()

    if request.method == "POST":
        type_ = request.form["type"]
        category = request.form["category"]
        amount = float(request.form["amount"])
        date = request.form["date"]
        description = request.form.get("description", "")

        cur.execute("""
            UPDATE transactions
            SET type = %s, category = %s, amount = %s, date = %s, description = %s
            WHERE id = %s
        """, (type_, category, amount, date, description, transaction_id))
        mysql.connection.commit()
        cur.close()
        flash('Transaction updated successfully!', 'success')
        return redirect(url_for("view_entries"))

    cur.execute("SELECT * FROM transactions WHERE id = %s", (transaction_id,))
    transaction = cur.fetchone()
    cur.close()

    if not transaction:
        flash('Transaction not found!', 'danger')
        return redirect(url_for("view_entries"))

    return render_template("edit_transaction.html", transaction=transaction)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config.get('ALLOWED_EXTENSIONS', set())

@app.route('/import', methods=["GET", "POST"])
@login_required
def import_transactions():
    if request.method == "POST":
        if 'file' not in request.files:
            flash('No file part', 'danger')
            return redirect(request.url)
        file = request.files['file']
        if file.filename == '':
            flash('No selected file', 'danger')
            return redirect(request.url)
        if file and allowed_file(file.filename):
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
            file.save(filepath)
            try:
                data = pd.read_excel(filepath)
                required_columns = ['Description', 'Date', 'Expense', 'Income', 'Closing Balance']
                if not all(col in data.columns for col in required_columns):
                    flash('Invalid file format. Ensure the headings are correct.', 'danger')
                    return redirect(request.url)
                cur = mysql.connection.cursor()
                for _, row in data.iterrows():
                    description = row['Description']
                    date = row['Date']
                    expense = row['Expense'] if not pd.isna(row['Expense']) else 0
                    income = row['Income'] if not pd.isna(row['Income']) else 0
                    if expense > 0:
                        type_ = 'Expense'
                        amount = expense
                    elif income > 0:
                        type_ = 'Income'
                        amount = income
                    else:
                        continue
                    cur.execute("""
                        INSERT INTO transactions (type, category, amount, date, description)
                        VALUES (%s, %s, %s, %s, %s)
                    """, (type_, 'Imported', amount, date, description))
                mysql.connection.commit()
                cur.close()
                flash('Transactions imported successfully!', 'success')
            except Exception as e:
                flash(f'Error processing file: {str(e)}', 'danger')
            finally:
                os.remove(filepath)
            return redirect(url_for('index'))
    return render_template('import.html')

# --------------------------------------------------------------------------
# User Authentication and Profile Routes
# --------------------------------------------------------------------------
@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        role = request.form['role']
        password = request.form['password']
        cursor = mysql.connection.cursor()
        cursor.execute(
            "SELECT id, username, role, password FROM staff WHERE username = %s AND role = %s", 
            (username, role)
        )
        user = cursor.fetchone()
        if user and user[3] == password:
            user_obj = User(id=user[0], username=user[1], role=user[2])
            login_user(user_obj)
            cursor.execute(
                "INSERT INTO attendance (staff_id, checkin_time) VALUES (%s, NOW())", 
                [user_obj.id]
            )
            mysql.connection.commit()
            flash('Login successful and checked in!', 'success')
            return redirect(url_for('index'))
        else:
            flash('Invalid credentials.', 'danger')
    return render_template('login.html')

@app.route('/logout')
@login_required
def logout():
    cursor = mysql.connection.cursor()
    cursor.execute("""
        UPDATE attendance 
        SET checkout_time = NOW() 
        WHERE staff_id = %s AND checkout_time IS NULL 
        ORDER BY checkin_time DESC LIMIT 1
    """, [current_user.id])
    mysql.connection.commit()
    logout_user()
    flash('Logged out successfully and checked out!', 'success')
    return redirect(url_for('login'))

@app.route('/profile', methods=['GET', 'POST'])
@login_required
def profile():
    cursor = mysql.connection.cursor()
    cursor.execute("""
        SELECT s.id, s.name, s.username, s.email, s.phone, s.position, 
               d.name AS department, r.name AS reportee, s.profile_picture
        FROM staff s
        LEFT JOIN department d ON s.department = d.id
        LEFT JOIN staff r ON s.reportee = r.id
        WHERE s.id = %s
    """, [current_user.id])
    user = cursor.fetchone()

    if not user:
        flash('User not found!', 'danger')
        return redirect(url_for('login'))

    user_data = {
        'id': user[0],
        'name': user[1],
        'username': user[2],
        'email': user[3],
        'phone': user[4],
        'position': user[5],
        'department': user[6],
        'reportee': user[7],
        'profile_picture': user[8]
    }

    if request.method == 'POST':
        name = request.form['name']
        phone = request.form['phone']
        profile_picture = None

        if 'profile_picture' in request.files:
            file = request.files['profile_picture']
            if file and file.filename != '':
                allowed_extensions = {'png', 'jpg', 'jpeg'}
                file_extension = file.filename.rsplit('.', 1)[-1].lower()
                if file_extension not in allowed_extensions:
                    flash('Invalid file type! Only PNG, JPG, and JPEG are allowed.', 'danger')
                    return redirect(url_for('profile'))
                filename = f"{current_user.id}_{secure_filename(file.filename)}"
                filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                file.save(filepath)
                profile_picture = f"uploads/{filename}"

        if profile_picture:
            cursor.execute("""
                UPDATE staff SET name = %s, phone = %s, profile_picture = %s WHERE id = %s
            """, (name, phone, profile_picture, current_user.id))
        else:
            cursor.execute("""
                UPDATE staff SET name = %s, phone = %s WHERE id = %s
            """, (name, phone, current_user.id))
        mysql.connection.commit()

        flash('Profile updated successfully!', 'success')
        return redirect(url_for('profile'))

    return render_template('profile.html', user=user_data)

# --------------------------------------------------------------------------
# Run the Flask App
# --------------------------------------------------------------------------
if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0')
